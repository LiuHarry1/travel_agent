"""统一文档加载器 - 支持多种格式并转换为 Markdown"""
from pathlib import Path
from typing import Dict, Any, List, Optional
import shutil
import uuid
import re
from urllib.parse import urlparse, urljoin

from .base import BaseLoader
from models.document import Document, DocumentType, DocumentStructure
from utils.exceptions import LoaderError
from utils.logger import get_logger

logger = get_logger(__name__)

# 按需导入
try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from bs4 import BeautifulSoup
    HAS_HTML = True
except ImportError:
    HAS_HTML = False

try:
    import markdownify
    HAS_MARKDOWNIFY = True
except ImportError:
    HAS_MARKDOWNIFY = False

try:
    from PIL import Image
    import io
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False

try:
    import zipfile
    HAS_ZIPFILE = True
except ImportError:
    HAS_ZIPFILE = False


class UnifiedLoader(BaseLoader):
    """统一文档加载器 - 将所有格式转换为 Markdown"""
    
    def __init__(self, static_dir: str = "static", base_url: str = ""):
        """
        Args:
            static_dir: 静态文件根目录，包含 sources/ 和 images/ 子目录
            base_url: 静态文件的基础 URL（如 http://localhost:8001）。如果为空，使用相对路径
        """
        self.static_dir = Path(static_dir)
        self.sources_dir = self.static_dir / "sources"
        self.images_dir = self.static_dir / "images"
        self.base_url = base_url.rstrip('/') if base_url else ""  # 移除末尾的斜杠
        
        # 确保目录存在
        self.sources_dir.mkdir(parents=True, exist_ok=True)
        self.images_dir.mkdir(parents=True, exist_ok=True)
        
        # 记录 base_url 配置（用于调试）
        if self.base_url:
            logger.info(f"UnifiedLoader initialized with base_url: '{self.base_url}' (applies to all file types: PDF, DOCX, HTML, Markdown)")
        else:
            logger.warning("UnifiedLoader initialized WITHOUT base_url (using relative paths). "
                         f"To enable full URLs, set STATIC_BASE_URL in .env file (e.g., STATIC_BASE_URL=http://localhost:8001)")
    
    def load(self, source: str, **kwargs) -> Document:
        """加载文档并转换为 Markdown"""
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {source}")
        
        doc_type = self._detect_type(path)
        
        # 保存原文件到 sources/
        original_filename = path.name
        file_id = kwargs.get("file_id") or self._generate_file_id(original_filename)
        saved_source_path = self._save_source_file(path, file_id, original_filename)
        
        logger.info(f"Loading {doc_type} file: {original_filename} (file_id: {file_id})")
        
        # 根据类型加载并转换
        structure = None
        try:
            if doc_type == DocumentType.PDF:
                content, structure = self._load_pdf(path, file_id)
            elif doc_type == DocumentType.DOCX:
                content, structure = self._load_docx(path, file_id)
            elif doc_type == DocumentType.HTML:
                content, structure = self._load_html(path, file_id)
            elif doc_type == DocumentType.TXT:
                content = self._load_txt(path)
            elif doc_type == DocumentType.MARKDOWN:
                content, structure = self._load_markdown(path, file_id)
            else:
                raise LoaderError(f"Unsupported document type: {doc_type}")
        except Exception as e:
            logger.error(f"Failed to load {doc_type} file: {str(e)}", exc_info=True)
            raise LoaderError(f"Failed to load {doc_type} file: {str(e)}") from e
        
        metadata = {
            "file_path": str(path),
            "file_name": original_filename,
            "file_size": path.stat().st_size,
            "saved_source_path": str(saved_source_path),
            "file_id": file_id,
            "original_type": doc_type.value,
            **kwargs.get("metadata", {})
        }
        
        # 添加页面信息到metadata（用于PDF）
        if structure and structure.total_pages:
            metadata["pages_info"] = structure.total_pages
        
        # 添加PDF元数据到metadata（title, author, subject, creator）
        if structure and structure.pdf_metadata:
            for key, value in structure.pdf_metadata.items():
                if value:  # 只添加非空值
                    metadata[f"pdf_{key}"] = value
        
        # Debug: Log metadata before creating Document
        logger.info(f"Document metadata before creation: {metadata}")
        logger.info(f"Document metadata keys: {list(metadata.keys())}")
        logger.info(f"Structure pdf_metadata: {structure.pdf_metadata if structure else 'None'}")
        
        return Document(
            content=content,
            source=str(saved_source_path),  # 使用保存后的路径
            doc_type=DocumentType.MARKDOWN,  # 统一为 Markdown
            metadata=metadata,
            structure=structure
        )
    
    def _detect_type(self, path: Path) -> DocumentType:
        """检测文档类型"""
        ext = path.suffix.lower()
        type_map = {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".doc": DocumentType.DOCX,
            ".html": DocumentType.HTML,
            ".htm": DocumentType.HTML,
            ".md": DocumentType.MARKDOWN,
            ".markdown": DocumentType.MARKDOWN,
            ".txt": DocumentType.TXT,
        }
        return type_map.get(ext, DocumentType.TXT)
    
    def _generate_file_id(self, filename: str) -> str:
        """生成唯一文件 ID"""
        return f"{uuid.uuid4().hex[:8]}_{Path(filename).stem}"
    
    def _save_source_file(self, source_path: Path, file_id: str, original_filename: str) -> Path:
        """保存原文件到 sources/ 目录"""
        # 保持原始扩展名
        ext = source_path.suffix
        saved_path = self.sources_dir / f"{file_id}{ext}"
        shutil.copy2(source_path, saved_path)
        logger.debug(f"Saved source file to: {saved_path}")
        return saved_path
    
    def _save_image(self, image_data: bytes, file_id: str, image_counter: int, ext: str = ".png") -> str:
        """保存图片并返回 URL（所有文件类型统一使用此方法，确保 base_url 正确应用）"""
        img_filename = f"{file_id}_image_{image_counter}{ext}"
        img_path = self.images_dir / img_filename
        
        with open(img_path, 'wb') as f:
            f.write(image_data)
        
        logger.debug(f"Saved image to: {img_path}")
        
        # 生成图片 URL（统一处理 base_url）
        relative_path = f"/static/images/{img_filename}"
        if self.base_url:
            # 使用配置的基础 URL（适用于所有文件类型：PDF、DOCX、HTML、Markdown）
            full_url = f"{self.base_url}{relative_path}"
            logger.info(f"Generated image URL with base_url ({self.base_url}): {full_url}")
            return full_url
        else:
            # 使用相对路径
            logger.warning(f"Generated image URL (relative, base_url is empty): {relative_path}. "
                         f"To use full URL, set STATIC_BASE_URL in .env file or config.")
            return relative_path
    
    def _load_pdf(self, path: Path, file_id: str) -> tuple[str, DocumentStructure]:
        """加载 PDF 并转换为 Markdown，返回内容和结构信息"""
        if not HAS_PDF:
            raise LoaderError("pdfplumber is required for PDF files. Install with: pip install pdfplumber")
        
        markdown_parts = []
        image_counter = 0
        pages_info = []
        pdf_metadata = {}
        tables_info = []
        
        try:
            with pdfplumber.open(path) as pdf:
                # 提取PDF元数据
                if pdf.metadata:
                    pdf_metadata = {
                        "title": pdf.metadata.get("Title", ""),
                        "author": pdf.metadata.get("Author", ""),
                        "subject": pdf.metadata.get("Subject", ""),
                        "creator": pdf.metadata.get("Creator", ""),
                    }
                    logger.info(f"Extracted PDF metadata: {pdf_metadata}")
                    logger.info(f"Raw PDF metadata from pdfplumber: {pdf.metadata}")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_start_char = len("\n".join(markdown_parts))
                    
                    # 提取文本
                    text = page.extract_text()
                    
                    # 提取表格
                    tables = page.extract_tables()
                    if tables:
                        for table_idx, table in enumerate(tables):
                            # 将表格转换为Markdown格式
                            table_md = self._table_to_markdown(table)
                            markdown_parts.append(f'\n<table index="{len(tables_info)}" page="{page_num}">\n{table_md}\n</table>\n\n')
                            tables_info.append({
                                "page": page_num,
                                "index": len(tables_info),
                                "rows": len(table),
                                "cols": len(table[0]) if table else 0
                            })
                    
                    if text:
                        # 使用 <page> 标签包裹每页内容，记录位置信息
                        markdown_parts.append(
                            f'<page page="{page_num}" start_char="{page_start_char}">\n\n'
                            f'## Page {page_num}\n\n{text}\n\n'
                        )
                    
                    # 尝试提取图片（使用 PyMuPDF/fitz 如果可用）
                    try:
                        import fitz  # PyMuPDF
                        pdf_doc = fitz.open(path)
                        page_obj = pdf_doc[page_num - 1]
                        
                        image_list = page_obj.get_images()
                        for img_idx, img in enumerate(image_list):
                            try:
                                xref = img[0]
                                base_image = pdf_doc.extract_image(xref)
                                image_bytes = base_image["image"]
                                image_ext = f".{base_image['ext']}"
                                
                                image_counter += 1
                                img_url = self._save_image(image_bytes, file_id, image_counter, image_ext)
                                markdown_parts.append(
                                    f'<img src="{img_url}" alt="Page {page_num} Image {img_idx + 1}" '
                                    f'page="{page_num}" image_index="{image_counter}" />\n\n'
                                )
                            except Exception as e:
                                logger.warning(f"Failed to extract image from PDF page {page_num}: {e}")
                                continue
                        
                        pdf_doc.close()
                    except ImportError:
                        # PyMuPDF 不可用，跳过图片提取
                        pass
                    except Exception as e:
                        logger.warning(f"Failed to extract images from PDF: {e}")
                    
                    # 关闭页面标签
                    if text:
                        page_end_char = len("\n".join(markdown_parts))
                        markdown_parts.append(f'</page>\n\n')
                        pages_info.append({
                            "page": page_num,
                            "start_char": page_start_char,
                            "end_char": page_end_char,
                            "bbox": list(page.bbox) if hasattr(page, 'bbox') else None
                        })
        except Exception as e:
            raise LoaderError(f"Failed to process PDF: {str(e)}") from e
        
        # 构建DocumentStructure
        structure = DocumentStructure(
            total_pages=len(pages_info),
            pdf_metadata=pdf_metadata if pdf_metadata else None,
            tables=tables_info if tables_info else None
        )
        
        return "\n".join(markdown_parts), structure
    
    def _table_to_markdown(self, table: list) -> str:
        """将表格转换为Markdown格式"""
        if not table or not table[0]:
            return ""
        
        md_lines = []
        # 表头
        header = table[0]
        md_lines.append("| " + " | ".join(str(cell) if cell else "" for cell in header) + " |")
        md_lines.append("| " + " | ".join("---" for _ in header) + " |")
        
        # 数据行
        for row in table[1:]:
            md_lines.append("| " + " | ".join(str(cell) if cell else "" for cell in row) + " |")
        
        return "\n".join(md_lines)
    
    def _load_docx(self, path: Path, file_id: str) -> tuple[str, DocumentStructure]:
        """加载 DOCX 并转换为 Markdown，提取图片"""
        if not HAS_DOCX:
            raise LoaderError("python-docx is required for DOCX files. Install with: pip install python-docx")
        
        if not HAS_ZIPFILE:
            raise LoaderError("zipfile module is required (should be built-in)")
        
        doc = DocxDocument(path)
        markdown_parts = []
        image_counter = 0
        
        # 处理段落
        for para in doc.paragraphs:
            if para.text.strip():
                markdown_parts.append(para.text)
        
        # 提取图片
        try:
            with zipfile.ZipFile(path, 'r') as docx_zip:
                # 查找图片文件
                image_files = [f for f in docx_zip.namelist() 
                              if f.startswith('word/media/')]
                
                for img_file in image_files:
                    try:
                        image_counter += 1
                        # 读取图片
                        img_data = docx_zip.read(img_file)
                        
                        # 确定图片扩展名
                        img_ext = Path(img_file).suffix
                        if not img_ext:
                            # 尝试从图片数据检测格式
                            if HAS_PIL:
                                try:
                                    img = Image.open(io.BytesIO(img_data))
                                    img_ext = f".{img.format.lower()}" if img.format else ".png"
                                except:
                                    img_ext = ".png"
                            else:
                                img_ext = ".png"
                        
                        # 保存图片
                        img_url = self._save_image(img_data, file_id, image_counter, img_ext)
                        
                        # 生成 HTML 图片链接
                        markdown_parts.append(f'\n<img src="{img_url}" alt="Image {image_counter}" />\n\n')
                    except Exception as e:
                        logger.warning(f"Failed to extract image {img_file} from DOCX: {e}")
                        continue
        except Exception as e:
            logger.warning(f"Failed to extract images from DOCX: {e}")
        
        # 构建DocumentStructure
        structure = DocumentStructure(
            total_sections=len([p for p in markdown_parts if p.strip()])
        )
        
        return "\n\n".join(markdown_parts), structure
    
    def _load_html(self, path: Path, file_id: str) -> tuple[str, DocumentStructure]:
        """加载 HTML 并转换为 Markdown，处理图片"""
        if not HAS_HTML:
            raise LoaderError("beautifulsoup4 is required for HTML files. Install with: pip install beautifulsoup4")
        
        try:
            with open(path, 'r', encoding='utf-8') as f:
                html_content = f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            with open(path, 'r', encoding='gbk', errors='ignore') as f:
                html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 提取并保存图片
        image_counter = 0
        for img_tag in soup.find_all('img'):
            src = img_tag.get('src', '')
            if not src:
                continue
            
            try:
                # 处理相对路径和绝对路径
                if src.startswith('http://') or src.startswith('https://'):
                    # 下载网络图片
                    if HAS_REQUESTS:
                        try:
                            response = requests.get(src, timeout=10, stream=True)
                            if response.status_code == 200:
                                img_data = response.content
                                # 从 Content-Type 或 URL 确定扩展名
                                content_type = response.headers.get('Content-Type', '')
                                if 'image/jpeg' in content_type or 'image/jpg' in content_type:
                                    img_ext = '.jpg'
                                elif 'image/png' in content_type:
                                    img_ext = '.png'
                                elif 'image/gif' in content_type:
                                    img_ext = '.gif'
                                else:
                                    img_ext = Path(urlparse(src).path).suffix or '.png'
                            else:
                                continue
                        except Exception as e:
                            logger.warning(f"Failed to download image {src}: {e}")
                            continue
                    else:
                        # 跳过网络图片（如果没有 requests）
                        continue
                else:
                    # 本地图片
                    # 处理相对路径
                    if src.startswith('/'):
                        # 绝对路径（相对于网站根目录）
                        img_path_abs = Path(path.parent / src.lstrip('/')).resolve()
                    else:
                        # 相对路径
                        img_path_abs = (path.parent / src).resolve()
                    
                    if img_path_abs.exists() and img_path_abs.is_file():
                        with open(img_path_abs, 'rb') as f:
                            img_data = f.read()
                        img_ext = img_path_abs.suffix or '.png'
                    else:
                        logger.warning(f"Image file not found: {img_path_abs}")
                        continue
                
                # 保存图片
                image_counter += 1
                img_url = self._save_image(img_data, file_id, image_counter, img_ext)
                
                # 替换为 HTML 图片标签
                alt_text = img_tag.get('alt', f'Image {image_counter}')
                img_tag.replace_with(f'<img src="{img_url}" alt="{alt_text}" />')
            except Exception as e:
                logger.warning(f"Failed to process image {src}: {e}")
                continue
        
        # 提取标题信息
        headings = []
        for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            level = int(heading.name[1])
            headings.append({
                "level": level,
                "text": heading.get_text().strip(),
                "id": heading.get('id', '')
            })
        
        # 转换为 Markdown
        if HAS_MARKDOWNIFY:
            markdown = markdownify.markdownify(str(soup), heading_style="ATX")
        else:
            # 简单转换：提取文本
            markdown = soup.get_text()
            logger.warning("markdownify not available, using simple text extraction")
        
        # 构建DocumentStructure
        structure = DocumentStructure(
            html_title=soup.title.string if soup.title else None,
            html_headings=headings if headings else None
        )
        
        return markdown, structure
    
    def _load_txt(self, path: Path) -> str:
        """加载纯文本文件"""
        try:
            return path.read_text(encoding='utf-8', errors='ignore')
        except Exception:
            # 尝试其他编码
            return path.read_text(encoding='gbk', errors='ignore')
    
    def _load_markdown(self, path: Path, file_id: str) -> tuple[str, DocumentStructure]:
        """加载 Markdown 文件，处理其中的图片"""
        try:
            content = path.read_text(encoding='utf-8')
        except UnicodeDecodeError:
            content = path.read_text(encoding='gbk', errors='ignore')
        
        # 查找 Markdown 中的图片链接
        # 匹配格式: ![alt](path) 或 <img src="path" />
        image_counter = 0
        
        # 处理 Markdown 图片语法 ![alt](path)
        def replace_markdown_image(match):
            nonlocal image_counter
            alt_text = match.group(1)
            img_path = match.group(2)
            
            # 如果是网络图片，保持原样或转换为 HTML
            if img_path.startswith('http://') or img_path.startswith('https://'):
                # 可选：下载网络图片
                if HAS_REQUESTS:
                    try:
                        response = requests.get(img_path, timeout=10)
                        if response.status_code == 200:
                            img_data = response.content
                            content_type = response.headers.get('Content-Type', '')
                            if 'image/jpeg' in content_type or 'image/jpg' in content_type:
                                img_ext = '.jpg'
                            elif 'image/png' in content_type:
                                img_ext = '.png'
                            elif 'image/gif' in content_type:
                                img_ext = '.gif'
                            else:
                                img_ext = Path(urlparse(img_path).path).suffix or '.png'
                            
                            image_counter += 1
                            img_url = self._save_image(img_data, file_id, image_counter, img_ext)
                            return f'<img src="{img_url}" alt="{alt_text}" />'
                    except Exception as e:
                        logger.warning(f"Failed to download image {img_path}: {e}")
                # 如果下载失败或没有 requests，保持原样
                return match.group(0)
            
            # 处理本地图片
            # 处理相对路径
            if img_path.startswith('/'):
                img_path_abs = Path(path.parent / img_path.lstrip('/')).resolve()
            else:
                img_path_abs = (path.parent / img_path).resolve()
            
            if img_path_abs.exists() and img_path_abs.is_file():
                image_counter += 1
                img_ext = img_path_abs.suffix or '.png'
                
                # 读取图片数据
                with open(img_path_abs, 'rb') as f:
                    img_data = f.read()
                
                # 保存图片
                img_url = self._save_image(img_data, file_id, image_counter, img_ext)
                
                # 转换为 HTML 图片标签
                return f'<img src="{img_url}" alt="{alt_text}" />'
            
            # 图片不存在，保持原样
            logger.warning(f"Image file not found: {img_path_abs}")
            return match.group(0)
        
        # 替换 Markdown 图片语法
        content = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', replace_markdown_image, content)
        
        # 处理 HTML img 标签中的相对路径
        def replace_html_image(match):
            nonlocal image_counter
            img_tag = match.group(0)
            src_match = re.search(r'src=["\']([^"\']+)["\']', img_tag)
            if not src_match:
                return img_tag
            
            src = src_match.group(1)
            
            # 如果是网络图片或已经是绝对路径，保持原样
            if src.startswith('http://') or src.startswith('https://') or src.startswith('/static/'):
                return img_tag
            
            # 处理本地相对路径
            if src.startswith('/'):
                img_path_abs = Path(path.parent / src.lstrip('/')).resolve()
            else:
                img_path_abs = (path.parent / src).resolve()
            
            if img_path_abs.exists() and img_path_abs.is_file():
                image_counter += 1
                img_ext = img_path_abs.suffix or '.png'
                
                with open(img_path_abs, 'rb') as f:
                    img_data = f.read()
                
                img_url = self._save_image(img_data, file_id, image_counter, img_ext)
                
                # 替换 src 属性
                return re.sub(r'src=["\'][^"\']+["\']', f'src="{img_url}"', img_tag)
            
            return img_tag
        
        # 替换 HTML img 标签
        content = re.sub(r'<img[^>]+>', replace_html_image, content)
        
        # 提取标题和代码块信息
        headings = []
        code_blocks = []
        
        # 提取Markdown标题
        heading_pattern = r'^(#{1,6})\s+(.+)$'
        for match in re.finditer(heading_pattern, content, re.MULTILINE):
            level = len(match.group(1))
            text = match.group(2).strip()
            headings.append({"level": level, "text": text})
        
        # 提取代码块
        code_block_pattern = r'```(\w+)?\n(.*?)```'
        for idx, match in enumerate(re.finditer(code_block_pattern, content, re.DOTALL)):
            lang = match.group(1) or ""
            code_blocks.append({
                "index": idx,
                "language": lang,
                "length": len(match.group(2))
            })
        
        # 构建DocumentStructure
        structure = DocumentStructure(
            md_headings=headings if headings else None,
            md_code_blocks=code_blocks if code_blocks else None
        )
        
        return content, structure
    
    def supports(self, doc_type: DocumentType) -> bool:
        """支持所有文档类型"""
        return True

